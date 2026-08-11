"""Render one stored CV into the .docx a recruiter receives.

Two substitutions happen, and keeping them apart is the whole design:

* **Content placeholders** -- the five in `TOKENS` -- are filled from the CV
  the model wrote. Three of them expand into many paragraphs, so they are
  replaced by *cloning* the placeholder's paragraph once per line: the clone
  keeps the direct run formatting (font, size, colour) that a template's
  paragraph style usually does not carry, which `insert_paragraph_before`
  would have dropped.
* **Locale IDs** -- any other `{KEY}` in the template -- are looked up in
  `cv_generator/l10n/values_<locale>.json`. Section headings use this, and so
  does anything else the owner of the template wants translated: writing
  ``{language_fr}, C1 - {LEVEL_FLUENT}`` in the Languages section renders as
  "Français, C1 - Courant" for a French posting and "French, C1 - Fluent" for
  an English one.

The consequence worth stating: nothing here reads the candidate's CV, guesses
a section format, or has an opinion about which languages a person speaks.
Whatever is not generated from the posting is written by whoever owns the
template, in their own words, once. `sections()` is pure and knows only the
model's JSON; `render_docx` knows only the template.

What the ATS sees is what this module protects: single column, no tables, no
text boxes, no images unless the caller asks, literal bullet glyphs rather
than a `numbering.xml` most templates do not ship, and headings in the
posting's own language so a parser recognises them.
"""

from __future__ import annotations

import copy
import io
import json
import re
import unicodedata
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.shared import Pt

from job_scraper import paths

#: A rendered line: the part set in bold, then the rest. Bold is the only
#: run-level decision any of these lines makes, so a tuple says it all.
Line = Tuple[str, str]

#: WordprocessingML namespace, for the one element python-docx does not wrap.
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

#: The placeholders filled from the generated CV. A template must carry every
#: one of them; everything else in it is the template owner's business.
TOKENS = (
    "CV_INTRODUCTION",
    "PROFILE_TEXT",
    "SKILLS_LIST_TEXT",
    "EXPERIENCES_PLACEHOLDER",
    "EDUCATION_PLACEHOLDER",
)

DEFAULT_LOCALE = "en"

#: Literal glyph, not a Word list. Templates rarely ship a `numbering.xml`, and
#: adding one to gain a bullet an extractor renders as nothing is a bad trade.
BULLET = "• "

#: Any `{KEY}` left in the document once the content tokens are expanded.
PLACEHOLDER_RE = re.compile(r"\{([A-Za-z_][A-Za-z0-9_]*)\}")


# LOADING

def load_l10n(locale: str) -> Dict[str, str]:
    """The locale file every `{KEY}` in a template is resolved against.

    Args:
        locale: Two-letter code as stored in `generated_cv.locale`.

    Returns:
        The locale's strings, falling back to English. `CV_SCHEMA` pins the
        four supported codes, but the column is TEXT, so an unknown code
        renders in English rather than raising.
    """
    path = paths.L10N / f"values_{locale}.json"

    if not path.exists():
        path = paths.L10N / f"values_{DEFAULT_LOCALE}.json"

    return json.loads(path.read_text(encoding="utf-8"))


# THE PURE CORE

def sections(cv: dict) -> Dict[str, List[Line]]:
    """Turn one stored CV into the lines that fill each content placeholder.

    No I/O, no locale, no Word: only the shape of the model's JSON.

    Args:
        cv: The `cv` object as stored in `generated_cv.cv`. Every key read here
            is required by `CV_SCHEMA` in `ai/cv_generation.py` -- the model
            call enforces the shape, so nothing is defended against. Note what
            is absent: languages are not in the schema and are not invented
            here either.

    Returns:
        One entry per token in `TOKENS`, each a list of lines to render as
        consecutive paragraphs.
    """
    blocks: Dict[str, List[Line]] = {
        "CV_INTRODUCTION": [("", cv["cv_introduction"])],
        "PROFILE_TEXT": [("", cv["profile_text"])],
        "SKILLS_LIST_TEXT": [
            (f"{group['competence']}: ", ", ".join(group["skills"]))
            for group in cv["skills"]
        ],
    }

    experiences: List[Line] = []

    for experience in cv["experiences"]:
        experiences.append((
            f"{experience['role']} - {experience['company']}",
            f", {experience['location']} "
            f"({experience['start_date']} - {experience['end_date']})",
        ))
        experiences += [("", BULLET + bullet)
                        for bullet in experience["bullets"]]

    blocks["EXPERIENCES_PLACEHOLDER"] = experiences

    education: List[Line] = []

    for entry in cv["education"]:
        education.append((
            f"{entry['diploma']} - {entry['institution']}",
            f", {entry['location']} ({entry['period']})",
        ))
        education.append(("", entry["degree"]))

        if entry["details"].strip():
            education.append(("", entry["details"]))

    blocks["EDUCATION_PLACEHOLDER"] = education

    return blocks


# THE DOCX ADAPTER

def render_docx(blocks: Dict[str, List[Line]], l10n: Dict[str, str],
                photo: bool,
                template: Path = paths.CV_TEMPLATE_DOCX) -> bytes:
    """Fill the template with `blocks`, resolve its locale IDs, return bytes.

    Args:
        blocks: Output of `sections`.
        l10n: Output of `load_l10n`, for every `{KEY}` the template carries
            that is not a content token.
        photo: Whether to keep any picture the template embeds. No default on
            purpose -- it is an ATS trade-off the caller must make out loud,
            not one this function makes behind the user's back.
        template: The .docx to fill.

    Returns:
        The document, as bytes -- what `st.download_button` wants, and what
        keeps this function out of the filesystem entirely.

    Raises:
        FileNotFoundError: If the template is not there. It is gitignored, so
            a fresh clone has only `CV_placeholder_example.docx`.
        RuntimeError: If the template is missing a content token, or carries a
            `{KEY}` that is neither a content token nor in the locale file --
            a typo that would otherwise be printed into the CV verbatim.
    """
    if not template.exists():
        raise FileNotFoundError(
            f"no CV template at {template} -- copy "
            f"user_info/CV_placeholder_example.docx to it and fill in your "
            f"name, contact details and languages")

    document = Document(str(template))
    placeholders = {paragraph.text.strip(): paragraph
                    for paragraph in document.paragraphs
                    if PLACEHOLDER_RE.fullmatch(paragraph.text.strip())}

    missing = [token for token in TOKENS if f"{{{token}}}" not in placeholders]

    if missing:
        raise RuntimeError(f"{template.name} is missing placeholders: "
                           f"{', '.join(missing)}")

    for token in TOKENS:
        _expand(placeholders[f"{{{token}}}"], blocks[token])

    _resolve(document, l10n, template.name)

    if not photo:
        _strip_drawings(document)

    _clean_text(document)
    _set_properties(document)

    stream = io.BytesIO()
    document.save(stream)

    return stream.getvalue()


def _expand(placeholder, lines: List[Line]) -> None:
    """Replace one placeholder paragraph with one paragraph per line.

    The clones carry the placeholder's own formatting, which is where a
    template's look actually lives.

    Args:
        placeholder: The paragraph whose text is the bare `{TOKEN}`.
        lines: What to put in its place; an empty list drops the paragraph.
    """
    element = placeholder._p

    for index, (bold, rest) in enumerate(lines):
        clone = copy.deepcopy(element)
        element.addprevious(clone)

        paragraph = placeholder.__class__(clone, placeholder._parent)
        _fill(paragraph, bold, rest)

        # Templates leave space after every body paragraph, which between
        # consecutive bullets turns a two-page CV into five. Only the last
        # line of the block keeps it, to separate the sections.
        if index < len(lines) - 1:
            paragraph.paragraph_format.space_after = Pt(0)

        # An experience header alone at the foot of a page reads as a mistake.
        paragraph.paragraph_format.keep_with_next = bool(bold)

    element.getparent().remove(element)


def _fill(paragraph, bold: str, rest: str) -> None:
    """Write one line into a cloned paragraph, bold lead-in first."""
    runs = paragraph.runs

    if not runs:
        return

    runs[0].text = bold or rest
    runs[0].bold = bool(bold)

    for spare in runs[1:]:
        spare._r.getparent().remove(spare._r)

    if bold and rest:
        # deepcopy of the run, not add_run(): a fresh run would inherit the
        # style defaults instead of the template's own formatting.
        tail = copy.deepcopy(runs[0]._r)
        runs[0]._r.addnext(tail)

        tail_run = paragraph.runs[1]
        tail_run.text = rest
        tail_run.bold = False


def _resolve(document, l10n: Dict[str, str], name: str) -> None:
    """Replace every remaining `{KEY}` with its string from the locale file.

    This is the seam that lets a template owner write their own Languages
    section -- or any other fixed text -- once, in IDs, and have it come out in
    the posting's language.

    Raises:
        RuntimeError: If a key is not in the locale file, listing the unknown
            ones. Silently leaving `{language_ja}` in a CV sent to a recruiter
            is the failure worth being loud about.
    """
    unknown = set()

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "{" not in run.text:
                continue

            unknown |= {key for key in PLACEHOLDER_RE.findall(run.text)
                        if key not in l10n}
            run.text = PLACEHOLDER_RE.sub(
                lambda match: l10n.get(match[1], match[0]), run.text)

    if unknown:
        raise RuntimeError(
            f"{name} uses locale IDs that no values_*.json defines: "
            f"{', '.join(sorted(unknown))} -- add them to every file in "
            f"cv_generator/l10n, or fix the spelling in the template")


def _strip_drawings(document) -> None:
    """Remove every embedded image from the body.

    Images are the most ATS-hostile thing a CV can carry: parsers skip them,
    and some reject the file outright. No-op on a template that has none.
    """
    for drawing in document.element.body.findall(f".//{{{_W}}}drawing"):
        run = drawing.getparent()
        run.getparent().remove(run)


#: Pictographs and the like. A pin emoji in front of a contact line is the
#: common case: several parsers turn it into mojibake and none of them read it.
DECORATION_RE = re.compile(r"[^\S ]|[☀-➿\U0001f000-\U0001faff]")


def _clean_text(document) -> None:
    """Strip decorative characters and the whitespace they leave behind.

    An emoji and the space after it are usually two separate runs, so removing
    the glyph alone would leave the line indented by one space.
    """
    for paragraph in document.paragraphs:
        dropped = False

        for run in paragraph.runs:
            cleaned = DECORATION_RE.sub("", run.text)
            dropped = dropped or cleaned != run.text
            run.text = cleaned

        if not dropped:
            continue

        # Left-strip across runs: the first one holding real text wins.
        for run in paragraph.runs:
            run.text = run.text.lstrip()

            if run.text:
                break


def _set_properties(document) -> None:
    """Name the document after the candidate, for the recruiter's file list."""
    name = candidate_name(document)
    document.core_properties.title = f"{name} - CV"
    document.core_properties.author = name


def candidate_name(document) -> str:
    """The candidate's name: the template's first non-empty paragraph."""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            return paragraph.text.strip()

    return "CV"


def filename(company: str, title: str,
             template: Path = paths.CV_TEMPLATE_DOCX) -> str:
    """A recruiter-legible filename for the generated CV.

    Args:
        company: The hiring company.
        title: The role as the posting names it.
        template: Where the candidate's name is read from.

    Returns:
        `Name_Company_Role.docx`, ASCII-only -- accented or non-Latin
        filenames still get mangled by enough upload forms to be worth
        avoiding.
    """
    name = candidate_name(Document(str(template))) if template.exists() else ""
    parts = [_slug(part) for part in (name, company, title)]
    stem = "_".join(part for part in parts if part)

    return f"{stem or 'cv'}.docx"


def _slug(text: str) -> str:
    """ASCII, underscores, nothing a file system or upload form dislikes."""
    folded = unicodedata.normalize("NFKD", text or "")
    folded = folded.encode("ascii", "ignore").decode()

    return re.sub(r"_+", "_", re.sub(r"[^A-Za-z0-9]+", "_", folded)).strip("_")
