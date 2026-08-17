"""The PDF renders without the docx toolkit's placeholder machinery, so its
own tests are lighter: it produces bytes at all, a photo is embedded when the
template has one, and it does not blow up on a template with none.
"""

import pytest

from cv_generator import docx_gen, pdf_gen
from job_scraper import paths
from tests.test_docx_gen import CV, TEMPLATE


@pytest.fixture
def blocks():
    return docx_gen.sections(CV)


@pytest.fixture
def l10n():
    return docx_gen.load_l10n("en")


def test_renders_bytes_from_the_example_template(blocks, l10n):
    data = pdf_gen.render_pdf(blocks, l10n, template=TEMPLATE)

    assert data.startswith(b"%PDF")


def test_a_missing_template_says_what_to_copy(tmp_path, blocks, l10n):
    with pytest.raises(FileNotFoundError, match="CV_placeholder_example"):
        pdf_gen.render_pdf(blocks, l10n, template=tmp_path / "gone.docx")


def test_contact_lines_skip_the_name_and_content_placeholders(blocks):
    from docx import Document

    lines = pdf_gen._contact_lines(Document(str(TEMPLATE)))

    assert lines == ["LOCATION   |   EMAIL  |  PHONE",
                     "LINKEDIN   |   PORTFOLIO"]


def test_a_template_with_no_photo_renders_none(blocks, l10n):
    from docx import Document

    assert pdf_gen._photo(Document(str(TEMPLATE))) is None


def test_languages_resolve_against_the_chosen_locale(blocks):
    from docx import Document

    lines = pdf_gen._resolved_languages(Document(str(TEMPLATE)),
                                        docx_gen.load_l10n("fr"))

    assert "Anglais, Langue maternelle" in lines
    assert "Français, B2 - Intermédiaire" in lines


@pytest.mark.skipif(not paths.CV_TEMPLATE_DOCX.exists(),
                    reason="real template is gitignored, not present here")
def test_the_real_template_s_photo_survives_into_the_pdf(blocks, l10n):
    data = pdf_gen.render_pdf(blocks, l10n, template=paths.CV_TEMPLATE_DOCX)

    assert len(data) > 10_000  # an embedded jpeg makes an empty-CV PDF this big
