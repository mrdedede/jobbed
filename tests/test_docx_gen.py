"""What the generated CV says, and that nothing is left unfilled.

`sections` is the pure half and gets pinned line by line. The rendering half
gets two assertions that matter more than the rest: no `{` survives anywhere in
the document, and a template naming an ID no locale file defines is refused.
Between them they cover the whole failure mode this module is shaped around --
a placeholder printed verbatim into a CV that goes to a recruiter. The shipped
example template really did name `{EDUCATION_TEXT}` where the real one names
`{EDUCATION_PLACEHOLDER}`.
"""

import io
import json
import zipfile

import pytest
from docx import Document

from cv_generator import docx_gen
from job_scraper import paths

TEMPLATE = paths.USER_INFO / "CV_placeholder_example.docx"

CV = {
    "cv_introduction": "Software Engineer, Lille",
    "profile_text": "Six years on distributed backends.",
    "skills": [
        {"competence": "Go", "skills": ["Go", "gRPC"]},
        {"competence": "Cloud", "skills": ["Kubernetes", "Terraform"]},
    ],
    "experiences": [
        {
            "role": "Backend Engineer", "company": "Acme",
            "location": "Lille, France", "start_date": "January, 2022",
            "end_date": "Present",
            "bullets": ["Cut p99 latency by 40%.", "Owned the release train."],
        },
        {
            "role": "Junior Engineer", "company": "Beta",
            "location": "Paris, France", "start_date": "January, 2020",
            "end_date": "December, 2021",
            "bullets": ["Built the onboarding flow."],
        },
    ],
    "education": [
        {
            "diploma": "MSc", "institution": "Centrale", "location": "Lille",
            "period": "2019 - 2021", "degree": "Computer Science",
            "details": "GPA 3.9",
        },
        {
            "diploma": "BSc", "institution": "USP", "location": "Sao Paulo",
            "period": "2015 - 2019", "degree": "Engineering",
            "details": "   ",
        },
    ],
}


@pytest.fixture
def french():
    return docx_gen.load_l10n("fr")


# sections

def test_a_skills_group_is_a_bold_label_and_a_joined_list():
    assert docx_gen.sections(CV)["SKILLS_LIST_TEXT"][0] == ("Go: ", "Go, gRPC")


def test_an_experience_is_a_header_then_one_line_per_bullet():
    lines = docx_gen.sections(CV)["EXPERIENCES_PLACEHOLDER"]

    assert lines[0] == ("Backend Engineer - Acme",
                        ", Lille, France (January, 2022 - Present)")
    assert lines[1:3] == [("", "• Cut p99 latency by 40%."),
                          ("", "• Owned the release train.")]
    assert lines[3] == ("Junior Engineer - Beta",
                        ", Paris, France (January, 2020 - December, 2021)")


def test_blank_education_details_add_no_line():
    lines = docx_gen.sections(CV)["EDUCATION_PLACEHOLDER"]

    # First entry: header, degree, details. Second: header, degree only.
    assert len(lines) == 5
    assert lines[2] == ("", "GPA 3.9")
    assert lines[3][0].startswith("BSc - USP")


def test_every_content_token_gets_a_block():
    assert set(docx_gen.sections(CV)) == set(docx_gen.TOKENS)


def test_sections_invents_no_languages():
    """Languages are the template owner's to write, not the model's to guess.

    They are absent from CV_SCHEMA on purpose, and nothing here may put them
    back by reading the candidate's CV or the database.
    """
    rendered = str(docx_gen.sections(CV))

    assert "language" not in rendered.lower()


# render_docx

def render(locale="fr", photo=False, template=TEMPLATE):
    """The sample CV rendered into a template, reopened."""
    data = docx_gen.render_docx(docx_gen.sections(CV),
                                docx_gen.load_l10n(locale), photo,
                                template=template)

    return Document(io.BytesIO(data))


@pytest.fixture
def rendered():
    return render()


def text_of(document):
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_no_placeholder_survives(rendered):
    assert "{" not in text_of(rendered) and "}" not in text_of(rendered)


def test_the_content_is_all_there(rendered):
    for expected in ("Backend Engineer - Acme", "• Cut p99 latency by 40%.",
                     "Go: Go, gRPC", "GPA 3.9", "Software Engineer, Lille"):
        assert expected in text_of(rendered)


def test_locale_ids_resolve_against_the_chosen_locale():
    """Headings and the template's own Languages lines, in one pass."""
    assert "EXPÉRIENCES PROFESSIONNELLES" in text_of(render("fr"))
    assert "PROFESSIONAL EXPERIENCES" in text_of(render("en"))


def test_the_templates_languages_are_written_in_the_postings_language():
    """The example template writes `{language_en}, {LEVEL_NATIVE}` itself.

    Nothing in the renderer knows which languages that is -- it only resolves
    the IDs the template happens to carry.
    """
    assert "English, Native" in text_of(render("en"))
    assert "Anglais, Langue maternelle" in text_of(render("fr"))
    assert "Français, B2 - Intermédiaire" in text_of(render("fr"))


def test_ids_resolve_mid_line_not_only_alone():
    """`{language_fr}, B2 - {LEVEL_INTERMEDIATE}` is two IDs in one run."""
    line = next(paragraph.text for paragraph in render("es").paragraphs
                if paragraph.text.startswith("Francés"))

    assert line == "Francés, B2 - Intermedio"


def test_clean_text_does_not_delete_a_run_that_holds_a_drawing(tmp_path):
    """`_clean_text` writes `run.text`, which clears a run's XML children --
    including a `w:drawing` a text run never held to begin with, if the run
    is touched without checking it actually carries a `w:t` first."""
    path = tmp_path / "with_drawing.docx"
    document = Document(str(TEMPLATE))
    paragraph = document.paragraphs[0]
    run = paragraph.add_run()
    drawing = run._r.makeelement(f"{{{docx_gen._W}}}drawing", {})
    run._r.append(drawing)
    document.save(str(path))

    reopened = Document(str(path))
    docx_gen._clean_text(reopened)

    assert reopened.element.body.findall(f".//{{{docx_gen._W}}}drawing")


def test_the_bold_lead_in_is_its_own_run(rendered):
    header = next(paragraph for paragraph in rendered.paragraphs
                  if paragraph.text.startswith("Backend Engineer"))

    assert [(run.text, run.bold) for run in header.runs] == [
        ("Backend Engineer - Acme", True),
        (", Lille, France (January, 2022 - Present)", False),
    ]


def test_a_job_s_last_bullet_keeps_space_before_the_next_job(rendered):
    """No gap within a job's bullets; a gap once the next line is a new job."""
    paragraphs = {paragraph.text: paragraph for paragraph in rendered.paragraphs}

    within_job = paragraphs["• Cut p99 latency by 40%."]
    last_bullet = paragraphs["• Owned the release train."]

    assert within_job.paragraph_format.space_after == 0
    assert last_bullet.paragraph_format.space_after != 0


def test_sections_stay_in_the_template_order(rendered):
    french = docx_gen.load_l10n("fr")
    lines = [paragraph.text for paragraph in rendered.paragraphs]
    order = [lines.index(french[key]) for key in
             ("PROFILE_TITLE", "SKILLS_TITLE", "EXPERIENCES_TITLE",
              "EDUCATION_TITLE", "LANGUAGES_TITLE")]

    assert order == sorted(order)


def broken_template(tmp_path, replace, with_):
    """A copy of the example template with one string swapped."""
    path = tmp_path / "broken.docx"
    source = zipfile.ZipFile(TEMPLATE)

    with zipfile.ZipFile(path, "w") as out:
        for item in source.infolist():
            data = source.read(item.filename)

            if item.filename == "word/document.xml":
                data = data.replace(replace, with_)

            out.writestr(item, data)

    return path


def test_a_template_missing_a_content_token_is_refused(tmp_path):
    path = broken_template(tmp_path, b"{EDUCATION_PLACEHOLDER}", b"")

    with pytest.raises(RuntimeError, match="EDUCATION_PLACEHOLDER"):
        render(template=path)


def test_a_template_naming_an_unknown_locale_id_is_refused(tmp_path):
    """A typo'd or undefined ID must not be printed into the CV verbatim."""
    path = broken_template(tmp_path, b"{language_en}", b"{language_ja}")

    with pytest.raises(RuntimeError, match="language_ja"):
        render(template=path)


def test_a_missing_template_says_what_to_copy(tmp_path):
    with pytest.raises(FileNotFoundError, match="CV_placeholder_example"):
        render(template=tmp_path / "gone.docx")


def test_an_unknown_locale_falls_back_to_english():
    assert docx_gen.load_l10n("de")["PROFILE_TITLE"] == "PROFILE"


def test_the_filename_is_ascii_and_names_the_role():
    name = docx_gen.filename("Açme & Co", "Ingénieur Back-end",
                             template=TEMPLATE)

    assert name == "Firstname_LASTNAME_Acme_Co_Ingenieur_Back_end.docx"
    assert name.isascii()


def test_the_locale_files_all_carry_the_same_keys():
    """An ID defined in one locale and forgotten in the others renders a CV
    that fails only for some postings."""
    files = sorted(paths.L10N.glob("values_*.json"))
    keys = [set(json.loads(path.read_text(encoding="utf-8"))) for path in files]

    assert len(files) == 4
    assert all(other == keys[0] for other in keys[1:])
