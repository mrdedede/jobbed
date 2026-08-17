"""Render one stored CV into the PDF a human reads.

The docx template still owns every word this module does not generate --
the candidate's name, contact line, photo and the Languages section's own
text -- so this reads them straight off the same template `docx_gen` fills,
rather than inventing a second place to store them. Only the four generated
sections (`sections()`'s output, from `docx_gen`) are laid out here.

No ATS trade-off applies to a PDF the way it does to a docx: a parser skips
the photo either way, so this module has no `photo` flag. The ATS-safe copy
stays `docx_gen.render_docx(..., photo=False)`.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (Image, Paragraph, SimpleDocTemplate, Spacer)

from cv_generator.docx_gen import Line, PLACEHOLDER_RE, _W, candidate_name
from job_scraper import paths

_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

_NAME_STYLE = ParagraphStyle("Name", fontName="Helvetica-Bold", fontSize=18,
                             spaceAfter=2 * mm)
_CONTACT_STYLE = ParagraphStyle("Contact", fontName="Helvetica", fontSize=9,
                                spaceAfter=6 * mm)
_HEADING_STYLE = ParagraphStyle("Heading", fontName="Helvetica-Bold",
                                fontSize=12, spaceBefore=4 * mm,
                                spaceAfter=2 * mm)
_BODY_STYLE = ParagraphStyle("Body", fontName="Helvetica", fontSize=10,
                             leading=13)


def render_pdf(blocks: Dict[str, List[Line]], l10n: Dict[str, str],
               template: Path = paths.CV_TEMPLATE_DOCX) -> bytes:
    """Fill `template`'s fixed content and `blocks` into a PDF, photo kept.

    Args:
        blocks: Output of `docx_gen.sections`.
        l10n: Output of `docx_gen.load_l10n`, for the section titles.
        template: The .docx whose name, contact line, photo and Languages
            text this reads -- never mutated.

    Returns:
        The PDF, as bytes -- what `st.download_button` wants.

    Raises:
        FileNotFoundError: If the template is not there.
    """
    if not template.exists():
        raise FileNotFoundError(
            f"no CV template at {template} -- copy "
            f"user_info/CV_placeholder_example.docx to it and fill in your "
            f"name, contact details and languages")

    document = Document(str(template))
    name = candidate_name(document)
    contact_lines = _contact_lines(document)
    photo = _photo(document)
    languages = _resolved_languages(document, l10n)

    stream = io.BytesIO()
    pdf = SimpleDocTemplate(stream, pagesize=A4,
                            leftMargin=20 * mm, rightMargin=20 * mm,
                            topMargin=18 * mm, bottomMargin=18 * mm)

    story = [Paragraph(name, _NAME_STYLE)]

    if photo is not None:
        photo.hAlign = "RIGHT"
        story.append(photo)

    for line in contact_lines:
        story.append(Paragraph(line, _CONTACT_STYLE))

    intro = blocks["CV_INTRODUCTION"][0][1]

    if intro:
        story.append(Paragraph(intro, _BODY_STYLE))

    story.append(Spacer(0, 4 * mm))

    for token, title_key in (("PROFILE_TEXT", "PROFILE_TITLE"),
                             ("SKILLS_LIST_TEXT", "SKILLS_TITLE"),
                             ("EXPERIENCES_PLACEHOLDER", "EXPERIENCES_TITLE"),
                             ("EDUCATION_PLACEHOLDER", "EDUCATION_TITLE")):
        story.append(Paragraph(l10n[title_key], _HEADING_STYLE))
        story += _flow(blocks[token])

    if languages:
        story.append(Paragraph(l10n["LANGUAGES_TITLE"], _HEADING_STYLE))
        story += [Paragraph(line, _BODY_STYLE) for line in languages]

    pdf.build(story)

    return stream.getvalue()


def _flow(lines: List[Line]) -> List:
    """One flowable per line, bold lead-in first, entries separated by a gap.

    Mirrors `docx_gen._expand`'s spacing rule: no gap between bullets of the
    same entry, a gap once the next line starts a new one (or this is the
    block's last line).
    """
    flowables = []

    for index, (bold, rest) in enumerate(lines):
        text = f"<b>{bold}</b>{rest}" if bold else rest
        is_last = index == len(lines) - 1
        ends_entry = is_last or bool(lines[index + 1][0])
        style = ParagraphStyle("Entry", parent=_BODY_STYLE,
                               spaceAfter=3 * mm if ends_entry else 0)
        flowables.append(Paragraph(text, style))

    return flowables


def _contact_lines(document) -> List[str]:
    """Every non-empty, non-placeholder paragraph before `{PROFILE_TITLE}`.

    The template owner may split contact details across more than one line
    (location/email/phone, then linkedin/portfolio), and may put
    `{CV_INTRODUCTION}` between the name and the contact block -- read
    generically, skipping placeholders, rather than assuming a fixed layout.
    """
    lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text == "{PROFILE_TITLE}":
            break

        if text and not PLACEHOLDER_RE.fullmatch(text):
            lines.append(text)

    return lines[1:]  # first non-empty line is the name, handled separately


def _photo(document) -> Optional[Image]:
    """The template's embedded picture, as a flowable -- `None` if it has none."""
    drawing = document.element.body.find(f".//{{{_W}}}drawing")

    if drawing is None:
        return None

    blip = drawing.find(f".//{{{_A}}}blip")
    part = document.part.related_parts[blip.get(f"{{{_R}}}embed")]
    width = 30 * mm

    return Image(io.BytesIO(part.blob), width=width, height=width)


def _resolved_languages(document, l10n: Dict[str, str]) -> List[str]:
    """The Languages section's own lines, locale IDs resolved.

    Every paragraph after `{LANGUAGES_TITLE}` is the template owner's text,
    the same `{language_fr}, B2 - {LEVEL_INTERMEDIATE}` pattern `docx_gen`
    resolves for the docx. Unknown IDs are left as-is here rather than
    raising: `render_docx` already validates the template on its own runs.
    """
    paragraphs = document.paragraphs
    start = next((i for i, p in enumerate(paragraphs)
                 if p.text.strip() == "{LANGUAGES_TITLE}"), None)

    if start is None:
        return []

    lines = []

    for paragraph in paragraphs[start + 1:]:
        text = paragraph.text.strip()

        if not text:
            continue

        lines.append(PLACEHOLDER_RE.sub(
            lambda match: l10n.get(match[1], match[0]), text))

    return lines
